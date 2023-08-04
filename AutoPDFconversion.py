import os
import shutil
from PyPDF2 import  PdfReader, PdfWriter
import PyPDF2
from docx import Document


def split_pdf(input_path, output_directory, chunk_size):
    # Create output directory if it doesn't exist
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    # Iterate over files in the input directory
    for filename in os.listdir(input_path):
        file_path = os.path.join(input_path, filename)

        if filename.endswith(".pdf"):
            # Process PDF files
            with open(file_path, 'rb') as input_file:
                pdf = PdfReader(input_file)
                total_pages = len(pdf.pages)
                num_chunks = total_pages // chunk_size
                if total_pages % chunk_size != 0:
                    num_chunks += 1

                for i in range(num_chunks):
                    start_page = i * chunk_size
                    end_page = min(start_page + chunk_size, total_pages)
                    output_pdf = PdfWriter()

                    for page in range(start_page, end_page):
                        output_pdf.add_page(pdf.pages[page])

                    output_file_path = os.path.join(output_directory, f'{filename}_chunk_{i+1}.pdf')
                    with open(output_file_path, 'wb') as output_file:
                        output_pdf.write(output_file)

                    print(f'Saved {output_file_path}')

        elif filename.endswith(".docx"):
            # Process Word documents
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            output_file_path = os.path.join(output_directory, f'{filename}.txt')

            with open(output_file_path, 'w', encoding='utf-8') as output_file:
                output_file.write(text)

            print(f'Saved {output_file_path}')

        else:
            print(f'Skipping {filename}: Not a PDF or Word document')


def convert_doc_to_txt(input_directory, output_directory):
    # Iterate over each Word document in the input directory
    for filename in os.listdir(input_directory):
        if filename.endswith(".docx"):
            doc_path = os.path.join(input_directory, filename)
            txt_path = os.path.join(output_directory, os.path.splitext(filename)[0] + ".txt")

            try:
                doc = Document(doc_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

                os.makedirs(os.path.dirname(txt_path), exist_ok=True)

                with open(txt_path, "w", encoding="utf-8") as txt_file:
                    txt_file.write(text)

                print(f"Successfully converted {filename} to {os.path.basename(txt_path)}")

            except Exception as e:
                print(f"Error converting {filename}: {e}")
                
def convert_pdf_to_txt(input_directory, output_directory):
    # Iterate over each PDF file in the input directory
    for filename in os.listdir(input_directory):
        if filename.endswith(".pdf"):
            pdf_path = os.path.join(input_directory, filename)
            txt_path = os.path.join(output_directory, os.path.splitext(filename)[0] + ".txt")

            # Convert PDF to TXT
            try:
                with open(pdf_path, "rb") as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    text = ""
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        text += page.extract_text()

                os.makedirs(os.path.dirname(txt_path), exist_ok=True)

                with open(txt_path, "w", encoding="utf-8") as txt_file:
                    txt_file.write(text)

                print(f"Successfully converted {filename} to {os.path.basename(txt_path)}")

            except Exception as e:
                print(f"Error converting {filename}: {e}")




def AutoConvertToTxt(input_directory, chunk_size, output_directory):
    # Create a temporary directory for storing the intermediate PDF chunks
    temp_directory = 'temp_chunks'
    os.makedirs(temp_directory, exist_ok=True)

    # Split PDF files
    split_pdf(input_directory, temp_directory, chunk_size)
    
    # Covert the split pdfs into text 
    convert_pdf_to_txt(temp_directory,output_directory)

    # Convert the remaining Word documents to TXT
    convert_doc_to_txt(input_directory, output_directory)

    # Delete the temporary directory
    shutil.rmtree(temp_directory)


# # Testing

# input_directory = "scrap"
# output_directory = "new_scrap"
# chunk_size = 2

# AutoConvertToTxt(input_directory, chunk_size, output_directory)
